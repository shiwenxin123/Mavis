"""
Marvis 文档深度分析 — 合同审查 / 运营数据分析 / 文案润色 / 图表生成。

可导入使用：
    from marvis.core.doc_analyzer import read_file, contract_review, data_analysis, generate_chart, polish_text
"""

import os
import re
import tempfile
from collections import Counter
from typing import Optional

from marvis.utils.fs import safe_read


# ─── 可选依赖 ───────────────────────────────────────────

_MISSING = {}

def _ensure(lib_name: str, pip_name: Optional[str] = None) -> bool:
    pn = pip_name or lib_name
    try:
        __import__(lib_name)
        return True
    except ImportError:
        _MISSING[pn] = f"python -m pip install {pn}"
        return False


# ─── 文件读取 ───────────────────────────────────────────

def read_file(filepath: str) -> Optional[str]:
    """读取文件内容，支持 PDF/DOCX/XLSX/CSV/TXT"""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        if not _ensure('pdfplumber'):
            return None
        import pdfplumber
        parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return '\n\n'.join(parts)

    elif ext in ('.docx', '.doc'):
        if not _ensure('docx', 'python-docx'):
            return None
        from docx import Document
        doc = Document(filepath)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(' | '.join(cells))
        return '\n\n'.join(parts)

    elif ext in ('.xlsx', '.xls'):
        if not _ensure('openpyxl'):
            return None
        import openpyxl
        wb = openpyxl.load_workbook(filepath, read_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"## Sheet: {sheet_name}")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [str(c) if c is not None else '' for c in row]
                parts.append(' | '.join(cells))
            if i > 500:
                parts.append('... (已截断)')
                break
        wb.close()
        return '\n'.join(parts)

    elif ext == '.csv':
        with open(filepath, 'r', encoding='utf-8-sig', errors='ignore') as f:
            return f.read()

    else:
        return safe_read(filepath)


# ─── 合同审查 ───────────────────────────────────────────

def contract_review(text: str) -> dict:
    """对合同文本进行审查，提取关键条款并识别风险点"""
    if not text:
        return {"error": "无法提取文本内容"}

    result: dict = {"extracted": {}, "risks": [], "checklist": []}

    # 提取关键信息
    amounts = re.findall(r'[\d,，]+\.?\d*\s*[万亿]?[元¥￥]', text)
    if amounts:
        result["extracted"]["金额"] = amounts[:10]

    dates = re.findall(r'\d{4}[/\-年]\d{1,2}[/\-月]\d{1,2}[日号]?', text)
    if dates:
        result["extracted"]["日期"] = list(set(dates))[:10]

    periods = re.findall(r'(\d+)\s*[个年月日天周][以之内前]', text)
    if periods:
        result["extracted"]["期限"] = periods[:10]

    penalty = re.findall(r'违约[金责].{0,50}', text)
    if penalty:
        result["extracted"]["违约条款"] = penalty[:5]

    confidentiality = re.findall(r'保密.{0,100}', text)
    if confidentiality:
        result["extracted"]["保密条款"] = confidentiality[:3]

    # 风险规则引擎
    risk_rules = [
        (r'不可抗力.{0,50}免除', "不可抗力免责条款过于宽泛"),
        (r'单方.{0,10}解除', "存在单方解除权，注意是否对等"),
        (r'违约金.{0,30}\d+%', "违约金可能过高，检查是否超过法律上限"),
        (r'概不负责|不承担任何', "免责条款，风险较高"),
        (r'最终解释权', "单方解释权条款，可能不公平"),
        (r'不可.{0,5}修改|不得变更', "条款不可修改，注意是否合理"),
        (r'自动.{0,5}续期|自动续费', "自动续期/续费条款，注意取消条件"),
        (r'仲裁.{0,30}指定', "仲裁机构由对方指定，可能不利"),
        (r'管辖.{0,30}甲方', "管辖权约定在甲方所在地，可能不利"),
        (r'预付款.{0,30}[5789]0%', "预付款比例过高，风险较大"),
        (r'验收.{0,30}视为合格', "默认验收合格条款，需注意异议期"),
    ]

    for pattern, desc in risk_rules:
        if re.search(pattern, text):
            result["risks"].append(desc)

    result["checklist"] = [
        "合同主体是否明确",
        "签约方是否有签约资格",
        "合同金额和付款方式是否清晰",
        "交付/服务标准是否明确",
        "验收标准和验收流程",
        "违约责任是否对等",
        "争议解决方式",
        "合同生效和终止条件",
        "保密条款是否合理",
        "知识产权归属",
    ]

    return result


# ─── 运营数据分析 ───────────────────────────────────────

def data_analysis(filepath: str) -> dict:
    """分析 Excel/CSV 数据：统计摘要、异常值、趋势"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in ('.xlsx', '.xls', '.csv'):
        return {"error": "仅支持 Excel/CSV 文件"}

    if not _ensure('pandas'):
        return {"error": "需要 pandas；运行 pip install pandas"}

    import pandas as pd

    try:
        if ext == '.csv':
            df = pd.read_csv(filepath, encoding='utf-8-sig', on_bad_lines='skip')
        else:
            df = pd.read_excel(filepath)
    except Exception as e:
        return {"error": f"读取文件失败: {e}"}

    result = {
        "shape": list(df.shape),
        "columns": list(df.columns),
        "dtypes": {str(col): str(dt) for col, dt in df.dtypes.items()},
        "head": df.head(5).to_dict(orient='records'),
        "null_counts": {str(k): int(v) for k, v in df.isnull().sum().to_dict().items()},
        "numeric_stats": {},
        "trends": [],
    }

    numeric_cols = df.select_dtypes(include='number').columns
    for col in numeric_cols:
        stats = df[col].describe().to_dict()
        result["numeric_stats"][str(col)] = {
            k: round(v, 2) if isinstance(v, float) else v
            for k, v in stats.items()
        }

        # 异常值（3σ）
        mean = df[col].mean()
        std = df[col].std()
        if std > 0:
            outliers = df[(df[col] < mean - 3 * std) | (df[col] > mean + 3 * std)]
            if len(outliers) > 0:
                result["trends"].append({
                    "type": "异常值",
                    "column": str(col),
                    "count": int(len(outliers)),
                    "message": f"列 '{col}' 有 {len(outliers)} 个异常值（超出3σ）"
                })

        # 趋势
        if len(df) >= 5:
            first_half = df[col].iloc[:len(df) // 2].mean()
            second_half = df[col].iloc[len(df) // 2:].mean()
            if first_half > 0 and second_half > 0:
                change = (second_half - first_half) / first_half * 100
                if abs(change) > 10:
                    direction = "上升" if change > 0 else "下降"
                    result["trends"].append({
                        "type": "趋势",
                        "column": str(col),
                        "change_pct": round(change, 1),
                        "message": f"列 '{col}' 整体{direction} {abs(round(change, 1))}%"
                    })

    # 缺失值
    for col, cnt in result["null_counts"].items():
        if cnt > 0:
            pct = cnt / len(df) * 100
            result["trends"].append({
                "type": "缺失值",
                "column": col,
                "count": int(cnt),
                "pct": round(pct, 1),
                "message": f"列 '{col}' 有 {int(cnt)} 个缺失值 ({round(pct, 1)}%)"
            })

    return result


# ─── 图表生成 ───────────────────────────────────────────

def generate_chart(
    filepath: str,
    chart_type: str = 'auto',
    output_path: Optional[str] = None,
    columns: Optional[list] = None,
) -> dict:
    """从 Excel/CSV 生成图表（line/bar/hist/pie/scatter/auto）"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in ('.xlsx', '.xls', '.csv'):
        return {"error": "仅支持 Excel/CSV 文件"}

    if not _ensure('pandas') or not _ensure('matplotlib'):
        return {"error": "需要 pandas 和 matplotlib"}

    import pandas as pd
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial']
    plt.rcParams['axes.unicode_minus'] = False

    try:
        if ext == '.csv':
            df = pd.read_csv(filepath, encoding='utf-8-sig', on_bad_lines='skip')
        else:
            df = pd.read_excel(filepath)

        numeric_cols = list(df.select_dtypes(include='number').columns)
        if not numeric_cols:
            return {"error": "没有数值列可用于生成图表"}

        if columns:
            numeric_cols = [c for c in columns if c in numeric_cols]

        if chart_type == 'auto':
            if len(numeric_cols) == 1:
                chart_type = 'hist'
            elif len(numeric_cols) <= 3:
                chart_type = 'line'
            else:
                chart_type = 'bar'

        fig, ax = plt.subplots(figsize=(12, 6))

        if chart_type == 'line':
            for col in numeric_cols[:5]:
                ax.plot(df.index, df[col], label=col, marker='o', markersize=3)
            ax.set_xlabel('Index')
            ax.legend()
        elif chart_type == 'bar':
            df[numeric_cols[:5]].plot.bar(ax=ax)
            ax.legend()
        elif chart_type == 'hist':
            df[numeric_cols[0]].plot.hist(ax=ax, bins=30)
        elif chart_type == 'pie':
            if len(numeric_cols) >= 1:
                col = numeric_cols[0]
                data = df[col].value_counts().head(10)
                ax.pie(data.values, labels=data.index, autopct='%1.1f%%')
        elif chart_type == 'scatter' and len(numeric_cols) >= 2:
            ax.scatter(df[numeric_cols[0]], df[numeric_cols[1]], alpha=0.5)
            ax.set_xlabel(numeric_cols[0])
            ax.set_ylabel(numeric_cols[1])

        ax.set_title(f'{os.path.basename(filepath)} - {chart_type}')
        plt.tight_layout()

        if not output_path:
            output_path = os.path.splitext(filepath)[0] + f'_{chart_type}.png'

        fig.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close(fig)

        return {"chart_path": os.path.abspath(output_path), "chart_type": chart_type}

    except Exception as e:
        return {"error": str(e)}


# ─── 文案润色 ───────────────────────────────────────────

def polish_text(text: str) -> dict:
    """分析文本质量，给出润色建议"""
    if not text:
        return {"error": "文本为空"}

    suggestions = []

    # 长句检测
    for i, sentence in enumerate(text.split('。')):
        if len(sentence) > 80:
            suggestions.append({
                "type": "长句拆分",
                "line": i + 1,
                "text": sentence[:100] + "..." if len(sentence) > 100 else sentence,
                "suggestion": "建议拆分为多个短句，提高可读性"
            })

    # 重复用词检测
    words = re.findall(r'[\u4e00-\u9fff]+', text)
    word_counts = Counter(words)
    for word, count in word_counts.most_common(20):
        if count > 3 and len(word) >= 2:
            suggestions.append({
                "type": "重复用词",
                "word": word,
                "count": count,
                "suggestion": f"'{word}' 出现 {count} 次，考虑使用同义词替换部分"
            })

    # 被动语态
    passive = re.findall(r'被.{2,10}[了着过]', text)
    if passive:
        suggestions.append({
            "type": "被动语态",
            "count": len(passive),
            "examples": passive[:3],
            "suggestion": "过多被动语态降低可读性，建议部分改为主动语态"
        })

    # 数字格式
    mixed_nums = re.findall(r'\d+\.?\d*[万亿]', text)
    if mixed_nums:
        suggestions.append({
            "type": "数字格式",
            "examples": mixed_nums[:3],
            "suggestion": "数字与单位间建议加空格，如 '1 万' 而非 '1万'"
        })

    return {
        "text_length": len(text),
        "suggestions": suggestions,
        "score": max(0, 100 - len(suggestions) * 5)
    }
