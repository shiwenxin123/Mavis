"""
Marvis 文件格式转换 — PDF/DOCX/XLSX/CSV/JSON/Markdown 互转。

可导入使用：
    from marvis.core.format_converter import pdf_to_text, xlsx_to_json, convert
"""

import csv
import json
import os
from typing import Any, Optional


def _ensure_lib(lib_name: str, pip_name: Optional[str] = None) -> bool:
    """检查可选依赖是否已安装"""
    try:
        __import__(lib_name)
        return True
    except ImportError:
        return False


def pdf_to_text(pdf_path: str) -> Optional[str]:
    """PDF 提取纯文本"""
    if not _ensure_lib("pdfplumber"):
        raise ImportError("需要安装 pdfplumber: pip install pdfplumber")
    import pdfplumber
    text_parts = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n\n".join(text_parts)


def docx_to_text(docx_path: str) -> Optional[str]:
    """DOCX 提取纯文本"""
    if not _ensure_lib("docx", "python-docx"):
        raise ImportError("需要安装 python-docx: pip install python-docx")
    from docx import Document
    doc = Document(docx_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def xlsx_to_json(xlsx_path: str) -> dict:
    """XLSX 转 JSON（每个 Sheet 为一个 key）"""
    if not _ensure_lib("openpyxl"):
        raise ImportError("需要安装 openpyxl: pip install openpyxl")
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    result = {}
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = []
        headers = None
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(c) if c else f"col_{j}" for j, c in enumerate(row)]
            else:
                rows.append(dict(zip(headers, [str(c) if c is not None else "" for c in row])))
        result[sheet] = rows
    wb.close()
    return result


def csv_to_json(csv_path: str) -> list[dict]:
    """CSV 转 JSON"""
    result = []
    with open(csv_path, "r", encoding="utf-8-sig", errors="ignore") as f:
        reader = csv.DictReader(f)
        for row in reader:
            result.append(dict(row))
    return result


def json_to_csv(data: Any, output_path: str):
    """JSON 转 CSV"""
    if isinstance(data, dict):
        for sheet_name, rows in data.items():
            if not rows:
                continue
            out_file = output_path.replace(".csv", f"_{sheet_name}.csv") if len(data) > 1 else output_path
            _write_csv(out_file, rows)
            print(f"已输出: {out_file}")
    elif isinstance(data, list) and data:
        _write_csv(output_path, data)
        print(f"已输出: {output_path}")


def _write_csv(path: str, rows: list[dict]):
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def xlsx_to_markdown(xlsx_path: str) -> str:
    """XLSX 转 Markdown 表格"""
    if not _ensure_lib("openpyxl"):
        raise ImportError("需要安装 openpyxl: pip install openpyxl")
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## {sheet_name}\n")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(c) if c is not None else "" for c in row]
            line = "| " + " | ".join(cells) + " |"
            parts.append(line)
            if i == 0:
                parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
        parts.append("")
    wb.close()
    return "\n".join(parts)


def convert(
    input_path: str,
    output_format: str = "json",
    output_path: Optional[str] = None,
) -> Any:
    """一站式格式转换入口"""
    ext = os.path.splitext(input_path)[1].lower()

    if ext == ".pdf":
        text = pdf_to_text(input_path)
        return {"text": text}
    elif ext in (".docx", ".doc"):
        text = docx_to_text(input_path)
        return {"text": text}
    elif ext == ".xlsx":
        if output_format == "markdown":
            return {"markdown": xlsx_to_markdown(input_path)}
        data = xlsx_to_json(input_path)
        if output_format == "csv" and output_path:
            json_to_csv(data, output_path)
        return data
    elif ext == ".csv":
        data = csv_to_json(input_path)
        return data
    elif ext == ".json":
        with open(input_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if output_format == "csv" and output_path:
            json_to_csv(data, output_path)
        return data
    else:
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            return {"text": f.read()}
