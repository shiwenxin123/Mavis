#!/usr/bin/env python3
"""Marvis 文档深度分析 - 合同审查/运营分析/文案润色/图表生成"""

import argparse
import json
import os
import sys


def ensure_lib(lib_name, pip_name=None):
    pip_name = pip_name or lib_name
    try:
        __import__(lib_name)
        return True
    except ImportError:
        print(f"需要安装依赖: python -m pip install {pip_name}", file=sys.stderr)
        return False


def read_file(filepath):
    """读取文件内容"""
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.pdf':
        if not ensure_lib('pdfplumber'):
            return None
        import pdfplumber
        parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return '\n\n'.join(parts)
    
    elif ext in ('.docx', '.doc'):
        if not ensure_lib('docx', 'python-docx'):
            return None
        from docx import Document
        doc = Document(filepath)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(' | '.join(cells))
        return '\n\n'.join(parts)
    
    elif ext in ('.xlsx', '.xls'):
        if not ensure_lib('openpyxl'):
            return None
        import openpyxl
        wb = openpyxl.load_workbook(filepath, read_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"## Sheet: {sheet_name}")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [str(c) if c is not None else '' for c in row]
                parts.append(' | '.join(cells))
        wb.close()
        return '\n'.join(parts)
    
    elif ext == '.csv':
        with open(filepath, 'r', encoding='utf-8-sig', errors='ignore') as f:
            return f.read()
    
    else:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


def contract_review(text):
    """合同审查 - 提取关键条款和风险点"""
    if not text:
        return {"error": "无法提取文本内容"}
    
    result = {
        "extracted": {},
        "risks": [],
        "checklist": []
    }
    
    import re
    
    # 提取合同金额
    amounts = re.findall(r'[\d,，]+\.?\d*\s*[万亿]?[元¥￥]', text)
    if amounts:
        result["extracted"]["金额"] = amounts
    
    # 提取日期
    dates = re.findall(r'\d{4}[\-/年]\d{1,2}[\-/月]\d{1,2}[日号]?', text)
    if dates:
        result["extracted"]["日期"] = list(set(dates))
    
    # 提取期限
    periods = re.findall(r'(\d+)\s*[个年月日天周][以之内前]', text)
    if periods:
        result["extracted"]["期限"] = periods
    
    # 提取违约金/赔偿
    penalty = re.findall(r'违约[金责].{0,50}', text)
    if penalty:
        result["extracted"]["违约条款"] = penalty[:5]
    
    # 提取保密条款
    confidentiality = re.findall(r'保密.{0,100}', text)
    if confidentiality:
        result["extracted"]["保密条款"] = confidentiality[:3]
    
    # 风险检测
    risk_rules = [
        (r'不可抗力.{0,50}免除', "不可抗力免责条款过于宽泛"),
        (r'单方.{0,10}解除', "存在单方解除权，注意是否对等"),
        (r'违约金.{0,30}\d+%', "违约金可能过高，检查是否超过法律上限"),
        (r'概不负责|不承担任何', "免责条款，风险较高"),
        (r'最终解释权', "单方解释权条款，可能不公平"),
        (r'不可.{0,5}修改|不得变更', "条款不可修改，注意是否合理"),
        (r'自动.{0,5}续期|自动续费', "自动续期/续费条款，注意取消条件"),
        (r'仲裁.{0,30}指定', "仲裁机构由对方指定，可能不利"),
        (r'管辖.{0,30}甲方', "管辖权约定在甲方所在地，可能不利"),
        (r'预付款.{0,30}[5789]0%', "预付款比例过高，风险较大"),
        (r'验收.{0,30}视为合格', "默认验收合格条款，需注意异议期"),
    ]
    
    for pattern, desc in risk_rules:
        if re.search(pattern, text):
            result["risks"].append(desc)
    
    # 合同审查检查项
    checklist_items = [
        "合同主体是否明确",
        "签约方是否有签约资格",
        "合同金额和付款方式是否清晰",
        "交付/服务标准是否明确",
        "验收标准和验收流程",
        "违约责任是否对等",
        "争议解决方式",
        "合同生效和终止条件",
        "保密条款是否合理",
        "知识产权归属",
    ]
    
    result["checklist"] = checklist_items
    
    return result


def data_analysis(filepath):
    """运营数据分析 - 分析 Excel/CSV 数据"""
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext not in ('.xlsx', '.xls', '.csv'):
        return {"error": "仅支持 Excel/CSV 文件"}
    
    try:
        if not ensure_lib('pandas'):
            return {"error": "需要 pandas"}
        import pandas as pd
        
        if ext == '.csv':
            df = pd.read_csv(filepath, encoding='utf-8-sig', on_bad_lines='skip')
        else:
            df = pd.read_excel(filepath)
        
        result = {
            "shape": list(df.shape),
            "columns": list(df.columns),
            "dtypes": {col: str(dt) for col, dt in df.dtypes.items()},
            "head": df.head(5).to_dict(orient='records'),
            "null_counts": df.isnull().sum().to_dict(),
            "numeric_stats": {},
            "trends": [],
        }
        
        # 数值列统计
        numeric_cols = df.select_dtypes(include='number').columns
        for col in numeric_cols:
            stats = df[col].describe().to_dict()
            result["numeric_stats"][col] = {k: round(v, 2) if isinstance(v, float) else v 
                                            for k, v in stats.items()}
            
            # 检测异常值（超过3个标准差）
            mean = df[col].mean()
            std = df[col].std()
            if std > 0:
                outliers = df[(df[col] < mean - 3*std) | (df[col] > mean + 3*std)]
                if len(outliers) > 0:
                    result["trends"].append({
                        "type": "异常值",
                        "column": col,
                        "count": len(outliers),
                        "message": f"列 '{col}' 有 {len(outliers)} 个异常值（超出3σ）"
                    })
            
            # 检测趋势（单调递增/递减）
            if len(df) >= 5:
                first_half = df[col].iloc[:len(df)//2].mean()
                second_half = df[col].iloc[len(df)//2:].mean()
                if first_half > 0 and second_half > 0:
                    change = (second_half - first_half) / first_half * 100
                    if abs(change) > 10:
                        direction = "上升" if change > 0 else "下降"
                        result["trends"].append({
                            "type": "趋势",
                            "column": col,
                            "change_pct": round(change, 1),
                            "message": f"列 '{col}' 整体{direction} {abs(round(change, 1))}%"
                        })
        
        # 缺失值检测
        for col, cnt in result["null_counts"].items():
            if cnt > 0:
                pct = cnt / len(df) * 100
                result["trends"].append({
                    "type": "缺失值",
                    "column": col,
                    "count": int(cnt),
                    "pct": round(pct, 1),
                    "message": f"列 '{col}' 有 {int(cnt)} 个缺失值 ({round(pct, 1)}%)"
                })
        
        return result
    
    except Exception as e:
        return {"error": str(e)}


def generate_chart(filepath, chart_type='auto', output_path=None, columns=None):
    """从 Excel/CSV 生成图表"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext not in ('.xlsx', '.xls', '.csv'):
        return {"error": "仅支持 Excel/CSV 文件"}
    
    try:
        if not ensure_lib('pandas'):
            return {"error": "需要 pandas"}
        if not ensure_lib('matplotlib'):
            return {"error": "需要 matplotlib"}
        
        import pandas as pd
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        
        # 中文字体支持
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial']
        plt.rcParams['axes.unicode_minus'] = False
        
        if ext == '.csv':
            df = pd.read_csv(filepath, encoding='utf-8-sig', on_bad_lines='skip')
        else:
            df = pd.read_excel(filepath)
        
        # 选择数值列
        numeric_cols = list(df.select_dtypes(include='number').columns)
        if not numeric_cols:
            return {"error": "没有数值列可用于生成图表"}
        
        if columns:
            numeric_cols = [c for c in columns if c in numeric_cols]
        
        # 自动选择图表类型
        if chart_type == 'auto':
            if len(numeric_cols) == 1:
                chart_type = 'hist'
            elif len(numeric_cols) <= 3:
                chart_type = 'line'
            else:
                chart_type = 'bar'
        
        fig, ax = plt.subplots(figsize=(12, 6))
        
        if chart_type == 'line':
            for col in numeric_cols[:5]:
                ax.plot(df.index, df[col], label=col, marker='o', markersize=3)
            ax.set_xlabel('Index')
            ax.legend()
        elif chart_type == 'bar':
            df[numeric_cols[:5]].plot.bar(ax=ax)
            ax.legend()
        elif chart_type == 'hist':
            df[numeric_cols[0]].plot.hist(ax=ax, bins=30)
        elif chart_type == 'pie':
            if len(numeric_cols) >= 1:
                col = numeric_cols[0]
                data = df[col].value_counts().head(10)
                ax.pie(data.values, labels=data.index, autopct='%1.1f%%')
        elif chart_type == 'scatter' and len(numeric_cols) >= 2:
            ax.scatter(df[numeric_cols[0]], df[numeric_cols[1]], alpha=0.5)
            ax.set_xlabel(numeric_cols[0])
            ax.set_ylabel(numeric_cols[1])
        
        ax.set_title(f'{os.path.basename(filepath)} - {chart_type}')
        plt.tight_layout()
        
        if not output_path:
            output_path = os.path.splitext(filepath)[0] + f'_{chart_type}.png'
        
        fig.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close(fig)
        
        return {"chart_path": os.path.abspath(output_path), "chart_type": chart_type}
    
    except Exception as e:
        return {"error": str(e)}


def polish_text(text):
    """文案润色建议"""
    if not text:
        return {"error": "文本为空"}
    
    suggestions = []
    
    # 检查长句
    for i, sentence in enumerate(text.split('。')):
        if len(sentence) > 80:
            suggestions.append({
                "type": "长句拆分",
                "line": i + 1,
                "text": sentence[:100] + "...",
                "suggestion": "建议拆分为多个短句，提高可读性"
            })
    
    # 检查重复用词
    import re
    words = re.findall(r'[\u4e00-\u9fff]+', text)
    from collections import Counter
    word_counts = Counter(words)
    for word, count in word_counts.most_common(20):
        if count > 3 and len(word) >= 2:
            suggestions.append({
                "type": "重复用词",
                "word": word,
                "count": count,
                "suggestion": f"'{word}' 出现 {count} 次，考虑使用同义词替换部分"
            })
    
    # 检查被动语态
    passive = re.findall(r'被.{2,10}[了着过]', text)
    if passive:
        suggestions.append({
            "type": "被动语态",
            "count": len(passive),
            "examples": passive[:3],
            "suggestion": "过多被动语态降低可读性，建议部分改为主动语态"
        })
    
    # 检查数字格式
    mixed_nums = re.findall(r'\d+\.?\d*[万亿]', text)
    if mixed_nums:
        suggestions.append({
            "type": "数字格式",
            "examples": mixed_nums[:3],
            "suggestion": "数字与单位间建议加空格，如 '1 万' 而非 '1万'"
        })
    
    return {
        "text_length": len(text),
        "suggestions": suggestions,
        "score": max(0, 100 - len(suggestions) * 5)
    }


def main():
    parser = argparse.ArgumentParser(description="Marvis 文档深度分析")
    parser.add_argument("filepath", help="文件路径")
    parser.add_argument("--action", choices=["contract", "analysis", "chart", "polish"], required=True, help="分析类型")
    parser.add_argument("--chart-type", choices=["line", "bar", "hist", "pie", "scatter", "auto"], default="auto", help="图表类型")
    parser.add_argument("--output", "-o", help="输出文件路径")
    parser.add_argument("--columns", help="指定列名（逗号分隔）")
    parser.add_argument("--json", action="store_true", help="JSON 格式输出")

    args = parser.parse_args()

    if not os.path.isfile(args.filepath):
        print(f"错误: 文件不存在 - {args.filepath}", file=sys.stderr)
        sys.exit(1)

    if args.action == "contract":
        text = read_file(args.filepath)
        result = contract_review(text)
        if args.json:
            print(json.dumps(result, ensure_ascii=False, indent=2))
        else:
            print("=" * 60)
            print("合同审查报告")
            print("=" * 60)
            if result.get("extracted"):
                print("\n📋 提取信息:")
                for key, val in result["extracted"].items():
                    print(f"  {key}: {val}")
            if result.get("risks"):
                print("\n⚠️ 风险提示:")
                for risk in result["risks"]:
                    print(f"  ⚡ {risk}")
            if result.get("checklist"):
                print("\n✅ 审查检查项:")
                for item in result["checklist"]:
                    print(f"  ☐ {item}")

    elif args.action == "analysis":
        result = data_analysis(args.filepath)
        if args.json:
            print(json.dumps(result, ensure_ascii=False, indent=2, default=str))
        else:
            print("=" * 60)
            print("运营数据分析")
            print("=" * 60)
            print(f"\n数据形状: {result.get('shape', [])}")
            print(f"列名: {result.get('columns', [])}")
            if result.get("numeric_stats"):
                print("\n📊 数值统计:")
                for col, stats in result["numeric_stats"].items():
                    print(f"  {col}: mean={stats.get('mean')}, std={stats.get('std')}")
            if result.get("trends"):
                print("\n📈 趋势与异常:")
                for t in result["trends"]:
                    print(f"  [{t['type']}] {t['message']}")

    elif args.action == "chart":
        cols = args.columns.split(',') if args.columns else None
        result = generate_chart(args.filepath, args.chart_type, args.output, cols)
        if args.json:
            print(json.dumps(result, ensure_ascii=False, indent=2))
        else:
            if "error" in result:
                print(f"生成失败: {result['error']}")
            else:
                print(f"✅ 图表已生成: {result['chart_path']}")
                print(f"   类型: {result['chart_type']}")

    elif args.action == "polish":
        text = read_file(args.filepath)
        if text:
            result = polish_text(text)
        else:
            result = {"error": "无法读取文件内容"}
        if args.json:
            print(json.dumps(result, ensure_ascii=False, indent=2))
        else:
            print("=" * 60)
            print(f"文案润色建议 (评分: {result.get('score', 'N/A')}/100)")
            print("=" * 60)
            if result.get("suggestions"):
                for s in result["suggestions"]:
                    print(f"\n  [{s['type']}]")
                    if s.get('text'):
                        print(f"    原文: {s['text'][:80]}")
                    print(f"    建议: {s['suggestion']}")
            else:
                print("\n  文案质量良好，无需调整。")


if __name__ == "__main__":
    main()