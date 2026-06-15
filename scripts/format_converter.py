#!/usr/bin/env python3
"""Marvis 文件格式转换 - 常见文件格式互转"""

import argparse
import json
import os
import sys


def ensure_lib(lib_name, pip_name=None):
    """确保依赖库已安装"""
    try:
        __import__(lib_name)
        return True
    except ImportError:
        pip_name = pip_name or lib_name
        print(f"需要安装依赖: pip install {pip_name}", file=sys.stderr)
        return False


def pdf_to_text(pdf_path):
    """PDF 提取文本"""
    if not ensure_lib('pdfplumber', 'pdfplumber'):
        return None
    import pdfplumber
    text_parts = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return '\n\n'.join(text_parts)


def docx_to_text(docx_path):
    """DOCX 提取文本"""
    if not ensure_lib('docx', 'python-docx'):
        return None
    from docx import Document
    doc = Document(docx_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(' | '.join(cells))
    return '\n\n'.join(parts)


def xlsx_to_json(xlsx_path):
    """XLSX 转 JSON"""
    if not ensure_lib('openpyxl', 'openpyxl'):
        return None
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    result = {}
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = []
        headers = None
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(c) if c else f'col_{j}' for j, c in enumerate(row)]
            else:
                rows.append(dict(zip(headers, [str(c) if c is not None else '' for c in row])))
        result[sheet] = rows
    wb.close()
    return result


def csv_to_json(csv_path):
    """CSV 转 JSON"""
    import csv
    result = []
    with open(csv_path, 'r', encoding='utf-8-sig', errors='ignore') as f:
        reader = csv.DictReader(f)
        for row in reader:
            result.append(dict(row))
    return result


def json_to_csv(data, output_path):
    """JSON 转 CSV"""
    import csv
    if isinstance(data, dict):
        for sheet_name, rows in data.items():
            if not rows:
                continue
            out_file = output_path.replace('.csv', f'_{sheet_name}.csv') if len(data) > 1 else output_path
            with open(out_file, 'w', encoding='utf-8-sig', newline='') as f:
                writer = csv.DictWriter(f, fieldnames=rows[0].keys())
                writer.writeheader()
                writer.writerows(rows)
            print(f"已输出: {out_file}")
    elif isinstance(data, list) and data:
        with open(output_path, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=data[0].keys())
            writer.writeheader()
            writer.writerows(data)
        print(f"已输出: {output_path}")


def xlsx_to_markdown(xlsx_path):
    """XLSX 转 Markdown 表格"""
    if not ensure_lib('openpyxl', 'openpyxl'):
        return None
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## {sheet_name}\n")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [str(c) if c is not None else '' for c in row]
            line = '| ' + ' | '.join(cells) + ' |'
            parts.append(line)
            if i == 0:
                parts.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        parts.append('')
    wb.close()
    return '\n'.join(parts)


def main():
    parser = argparse.ArgumentParser(description='Marvis 文件格式转换')
    parser.add_argument('input', help='输入文件路径')
    parser.add_argument('--output', '-o', help='输出文件路径')
    parser.add_argument('--format', '-f', choices=['text', 'markdown', 'json', 'csv', 'html'], help='输出格式')

    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"错误: 文件不存在 - {args.input}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(args.input)[1].lower()
    
    if ext == '.pdf':
        text = pdf_to_text(args.input)
        if text is None:
            sys.exit(1)
        if args.format == 'json' or args.output and args.output.endswith('.json'):
            print(json.dumps({'text': text}, ensure_ascii=False, indent=2))
        else:
            print(text)

    elif ext in ('.docx', '.doc'):
        text = docx_to_text(args.input)
        if text is None:
            sys.exit(1)
        if args.format == 'json' or args.output and args.output.endswith('.json'):
            print(json.dumps({'text': text}, ensure_ascii=False, indent=2))
        else:
            print(text)

    elif ext == '.xlsx':
        fmt = args.format or 'json'
        if fmt == 'markdown':
            md = xlsx_to_markdown(args.input)
            if md:
                print(md)
        else:
            data = xlsx_to_json(args.input)
            if data is None:
                sys.exit(1)
            output = args.output
            if fmt == 'csv' or (output and output.endswith('.csv')):
                json_to_csv(data, output or 'output.csv')
            else:
                print(json.dumps(data, ensure_ascii=False, indent=2))

    elif ext == '.csv':
        data = csv_to_json(args.input)
        fmt = args.format or 'json'
        output = args.output
        if fmt == 'csv' or (output and output.endswith('.csv')):
            # CSV to CSV (maybe with different delimiter)
            print(json.dumps(data, ensure_ascii=False, indent=2))
        else:
            print(json.dumps(data, ensure_ascii=False, indent=2))

    elif ext == '.json':
        with open(args.input, 'r', encoding='utf-8') as f:
            data = json.load(f)
        output = args.output
        if args.format == 'csv' or (output and output.endswith('.csv')):
            json_to_csv(data, output or 'output.csv')
        else:
            print(json.dumps(data, ensure_ascii=False, indent=2))

    else:
        # 普通文本文件，直接输出
        with open(args.input, 'r', encoding='utf-8', errors='ignore') as f:
            print(f.read())


if __name__ == '__main__':
    main()
